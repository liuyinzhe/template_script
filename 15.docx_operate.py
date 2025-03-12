from docx import Document


def read_docx(doc_file):
    doc = Document(str(doc_file))
    #输出每一段/行的内容
    for para in doc.paragraphs:
        line = str(para.text).strip()
        print(line)


##################### 高亮  ########

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import re

def write_text_with_highlight(doc, text):
    # 分割普通文本和高亮内容
    parts = re.split(r'(\[.*?\])', text)
    paragraph = doc.add_paragraph()
    for part in parts:
        if not part:  # 跳过空字符串
            continue
        # 如果是方括号包裹的内容，设置高亮
        if part.startswith('[') and part.endswith(']'):
            run = paragraph.add_run(part)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            '''
            WD_COLOR_INDEX.RED
            WD_COLOR_INDEX.GREEN
            枚举值	颜色效果	备注
            WD_COLOR_INDEX.AUTO	自动（通常无色）	默认值，无高亮效果
            WD_COLOR_INDEX.BLACK	黑色高亮	
            WD_COLOR_INDEX.BLUE	蓝色高亮	
            WD_COLOR_INDEX.BRIGHT_GREEN	亮绿色高亮	
            WD_COLOR_INDEX.DARK_BLUE	深蓝色高亮	
            WD_COLOR_INDEX.DARK_RED	深红色高亮	注意拼写是 DARK_RED（文档可能拼写为 DARK_RED）
            WD_COLOR_INDEX.DARK_YELLOW	暗黄色高亮	
            WD_COLOR_INDEX.GRAY_25	25% 灰度高亮	浅灰色
            WD_COLOR_INDEX.GRAY_50	50% 灰度高亮	中灰色
            WD_COLOR_INDEX.GREEN	绿色高亮	
            WD_COLOR_INDEX.PINK	粉色高亮	
            WD_COLOR_INDEX.RED	红色高亮	
            WD_COLOR_INDEX.TEAL	青色高亮	类似蓝绿色
            WD_COLOR_INDEX.TURQUOISE	绿松石色高亮	
            WD_COLOR_INDEX.VIOLET	紫色高亮	
            WD_COLOR_INDEX.WHITE	白色高亮	可能显示为无效果（与背景同色）
            WD_COLOR_INDEX.YELLOW	黄色高亮	注意拼写是 YELLOW
            '''
        else:
            paragraph.add_run(part)

# 示例文本
sample_text = "这是一个示例文本，其中包含[需要高亮]的部分，以及另一个[高亮内容]。"

# 创建Word文档并写入内容
doc = Document()
write_text_with_highlight(doc, sample_text)
doc.save("highlighted_content.docx")

print("Word文档已生成，高亮内容已标记。")
